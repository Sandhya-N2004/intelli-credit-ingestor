# ==========================================================
# INTELLI-CREDIT : PART 3 - RECOMMENDATION ENGINE
# ==========================================================

import json
import numpy as np
import shap
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document


# ==========================================================
# LOAD JSON
# ==========================================================

def load_json(file):

    with open(file) as f:
        return json.load(f)


# ==========================================================
# FIVE C's CREDIT MODEL
# ==========================================================

def compute_five_cs(part1, part2):

    financials = part1["financials"]
    fraud = part1["fraud_analysis"]

    revenue = financials.get("annual_revenue", 0)
    profit = financials.get("net_profit", 0)

    fraud_score = fraud.get("fraud_score", 0)

    risk_score = part2["final_risk"]

    # Character
    character = 1 - min(risk_score, 1)

    # Capacity (profitability / repayment ability)
    if revenue:
        capacity = max(0, min(1, profit / revenue))
    else:
        capacity = 0

    # Capital (proxy for financial strength)
    capital = min(1, revenue / 500000)

    # Collateral
    collateral = max(0.2, 1 - fraud_score)

    # Conditions (sector + macro risk)
    conditions = 1 - (
        part2["sector_risk"] +
        part2["news_risk"]
    ) / 2

    five_cs = {

        "Character": round(character, 3),
        "Capacity": round(capacity, 3),
        "Capital": round(capital, 3),
        "Collateral": round(collateral, 3),
        "Conditions": round(conditions, 3)

    }

    return five_cs


# ==========================================================
# CREDIT SCORE
# ==========================================================

def compute_credit_score(five_cs):

    weights = {

        "Character": 0.30,
        "Capacity": 0.25,
        "Capital": 0.20,
        "Collateral": 0.15,
        "Conditions": 0.10

    }

    score = 0

    for k, v in five_cs.items():
        score += weights[k] * v

    return round(score, 3)


# ==========================================================
# LOAN DECISION ENGINE
# ==========================================================

def loan_decision(score, revenue):

    if score < 0.4:

        decision = "REJECTED"
        loan = 0
        rate = None
        reason = "High credit risk detected"

    elif score < 0.65:

        decision = "CONDITIONAL APPROVAL"
        loan = revenue * 0.05
        rate = 13.5
        reason = "Moderate risk profile"

    else:

        decision = "APPROVED"
        loan = revenue * 0.10
        rate = 10.5
        reason = "Strong financial health"

    return {

        "decision": decision,
        "loan": round(loan, 2),
        "rate": rate,
        "reason": reason

    }


# ==========================================================
# EXPLAINABLE AI (SHAP)
# ==========================================================

def explain_model(five_cs):

    features = np.array(list(five_cs.values())).reshape(1, -1)

    background = np.random.rand(50, 5)

    explainer = shap.KernelExplainer(lambda x: np.sum(x, axis=1), background)

    shap_values = explainer.shap_values(features)

    feature_names = list(five_cs.keys())

    contributions = dict(zip(feature_names, shap_values[0]))

    print("\nExplainable AI Contribution:")

    for k, v in contributions.items():
        print(k, ":", round(v, 3))

    return contributions


# ==========================================================
# DONUT CHART
# ==========================================================

def donut_chart(five_cs, company):

    labels = list(five_cs.keys())

    values = [max(0, v) for v in five_cs.values()]

    fig, ax = plt.subplots()

    wedges, texts = ax.pie(

        values,
        labels=labels,
        startangle=90,
        wedgeprops=dict(width=0.4)

    )

    centre_circle = plt.Circle((0, 0), 0.70, fc='white')

    fig.gca().add_artist(centre_circle)

    plt.text(0, 0, "Credit\nProfile", ha='center', va='center', fontsize=12)

    plt.title("Five C's Credit Profile")

    file = f"{company}_credit_donut.png"

    plt.savefig(file)

    print("Donut chart saved:", file)


# ==========================================================
# RISK BREAKDOWN PIE CHART
# ==========================================================

def risk_breakdown_chart(part2, company):

    labels = [

        "Financial Risk",
        "Fraud Risk",
        "News Risk",
        "Promoter Risk",
        "Sector Risk",
        "Litigation Risk"

    ]

    values = [

        abs(part2["financial_risk"]),
        abs(part2["fraud_risk"]),
        abs(part2["news_risk"]),
        abs(part2["promoter_risk"]),
        abs(part2["sector_risk"]),
        abs(part2["litigation_risk"])

    ]

    plt.figure()

    plt.pie(values, labels=labels, autopct='%1.1f%%')

    plt.title("Risk Driver Analysis")

    file = f"{company}_risk_breakdown.png"

    plt.savefig(file)

    print("Risk breakdown chart saved:", file)


# ==========================================================
# MONTE CARLO DEFAULT SIMULATION
# ==========================================================

def simulate_default_probability(score):

    simulations = 10000

    risk = []

    for i in range(simulations):

        shock = np.random.normal(0, 0.05)

        simulated = score + shock

        if simulated < 0.5:
            risk.append(1)
        else:
            risk.append(0)

    probability = np.mean(risk)

    print("Default Probability:", round(probability, 3))

    return probability


# ==========================================================
# SECTOR BENCHMARK
# ==========================================================

def sector_benchmark(score):

    sector_avg = 0.55

    if score > sector_avg:
        status = "Above sector average"
    else:
        status = "Below sector average"

    return {

        "sector_avg": sector_avg,
        "comparison": status

    }


# ==========================================================
# CAM GENERATOR
# ==========================================================

def generate_cam(company, five_cs, decision, default_prob, benchmark):

    doc = Document()

    doc.add_heading("Credit Appraisal Memo", 0)

    doc.add_paragraph("Company: " + company)

    doc.add_paragraph("Date: " + str(datetime.now()))

    doc.add_heading("Five C's Analysis", 1)

    for k, v in five_cs.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Loan Decision", 1)

    doc.add_paragraph("Decision: " + decision["decision"])
    doc.add_paragraph("Recommended Loan: ₹" + str(decision["loan"]))
    doc.add_paragraph("Interest Rate: " + str(decision["rate"]))
    doc.add_paragraph("Reason: " + decision["reason"])

    doc.add_heading("Risk Simulation", 1)

    doc.add_paragraph("Default Probability: " + str(round(default_prob, 3)))

    doc.add_heading("Sector Benchmark", 1)

    doc.add_paragraph("Sector Average Score: " + str(benchmark["sector_avg"]))
    doc.add_paragraph("Comparison: " + benchmark["comparison"])

    file = f"CAM_{company}.docx"

    doc.save(file)

    print("CAM generated:", file)


# ==========================================================
# AI DECISION EXPLANATION
# ==========================================================

def explain_decision(five_cs, decision, default_prob, company):

    explanation = []

    explanation.append("AI CREDIT DECISION EXPLANATION\n")

    explanation.append(f"Loan Decision: {decision['decision']}")
    explanation.append(f"Recommended Loan: ₹{decision['loan']}")
    explanation.append(f"Interest Rate: {decision['rate']}")

    explanation.append("\nFive C's Evaluation:")

    for k, v in five_cs.items():

        if v < 0.4:
            explanation.append(f"{k} risk is HIGH ({v})")

        elif v < 0.7:
            explanation.append(f"{k} risk is MODERATE ({v})")

        else:
            explanation.append(f"{k} is STRONG ({v})")

    explanation.append(f"\nSimulated Default Probability: {round(default_prob, 3)}")

    file = f"{company}_decision_explanation.txt"

    with open(file, "w", encoding="utf-8") as f:
        f.write("\n".join(explanation))

    print("AI explanation saved:", file)


# ==========================================================
# MAIN PIPELINE
# ==========================================================

def run(part1_file, part2_file):

    part1 = load_json(part1_file)

    part2 = load_json(part2_file)

    company = part1["company"]

    print("\nRunning Recommendation Engine for:", company)

    five_cs = compute_five_cs(part1, part2)

    score = compute_credit_score(five_cs)

    revenue = part1["financials"].get("annual_revenue", 0)

    decision = loan_decision(score, revenue)

    explain_model(five_cs)

    donut_chart(five_cs, company)

    risk_breakdown_chart(part2, company)

    default_prob = simulate_default_probability(score)

    benchmark = sector_benchmark(score)

    generate_cam(company, five_cs, decision, default_prob, benchmark)

    explain_decision(five_cs, decision, default_prob, company)

    print("\nFINAL CREDIT SCORE:", score)

    print("DECISION:", decision)


# ==========================================================
# EXECUTION
# ==========================================================

if __name__ == "__main__":

    run(

        "output_Tata_Steel.json",
        "part2_output.json"

    )